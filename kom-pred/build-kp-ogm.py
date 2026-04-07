"""Build KP for granulator OGM-1.5 from original template. v3"""
import copy
from docx import Document
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.shared import Cm

doc = Document("КОММЕРЧЕСКОЕ ПРЕДЛОЖЕНИЕ макет.docx")
body = doc.element.body


# ─── Low-level XML helpers (defined early, used throughout) ───

def make_run(text, bold=False, size=22, color="1A1A1A", font="Arial"):
    r = OxmlElement("w:r")
    rPr = OxmlElement("w:rPr")
    rFonts = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs"): rFonts.set(qn(a), font)
    rPr.append(rFonts)
    for tag, val in [("w:sz", str(size)), ("w:szCs", str(size))]:
        el = OxmlElement(tag); el.set(qn("w:val"), val); rPr.append(el)
    if bold: rPr.append(OxmlElement("w:b"))
    clr = OxmlElement("w:color"); clr.set(qn("w:val"), color); rPr.append(clr)
    r.append(rPr)
    t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = text; r.append(t)
    return r


def make_para(text, bold=False, size=22, color="1A1A1A", space_before=0, space_after=0, indent_left=0):
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    if space_before or space_after:
        sp = OxmlElement("w:spacing")
        if space_before: sp.set(qn("w:before"), str(space_before))
        if space_after:  sp.set(qn("w:after"),  str(space_after))
        pPr.append(sp)
    if indent_left:
        ind = OxmlElement("w:ind")
        ind.set(qn("w:left"), str(indent_left))
        pPr.append(ind)
    p.append(pPr)
    p.append(make_run(text, bold=bold, size=size, color=color))
    return p


def make_cell(text, bold=False, width=None, shading=None, size=20):
    tc = OxmlElement("w:tc")
    tcPr = OxmlElement("w:tcPr")
    if width:
        tcW = OxmlElement("w:tcW"); tcW.set(qn("w:w"), str(width)); tcW.set(qn("w:type"), "dxa"); tcPr.append(tcW)
    tcBorders = OxmlElement("w:tcBorders")
    for side in ("top", "bottom", "left", "right"):
        b = OxmlElement(f"w:{side}")
        b.set(qn("w:val"), "single"); b.set(qn("w:sz"), "4")
        b.set(qn("w:color"), "CCCCCC"); b.set(qn("w:space"), "0")
        tcBorders.append(b)
    tcPr.append(tcBorders)
    if shading:
        shd = OxmlElement("w:shd"); shd.set(qn("w:val"), "clear"); shd.set(qn("w:fill"), shading); tcPr.append(shd)
    tcMar = OxmlElement("w:tcMar")
    for side, val in [("top","60"),("bottom","60"),("start","100"),("end","100")]:
        m = OxmlElement(f"w:{side}"); m.set(qn("w:w"), val); m.set(qn("w:type"), "dxa"); tcMar.append(m)
    tcPr.append(tcMar)
    tc.append(tcPr)
    tc.append(make_para(text, bold=bold, size=size))
    return tc


CONTENT_WIDTH = 11198  # A4, margins: left=284, right=424 → 11906-284-424

def apply_tbl_indent(tbl_el, indent_dxa, right_indent_dxa=None):
    """Add left indent + shrink table width so right margin is also respected."""
    if right_indent_dxa is None:
        right_indent_dxa = indent_dxa
    target_width = CONTENT_WIDTH - indent_dxa - right_indent_dxa

    tblPr = tbl_el.find(qn("w:tblPr"))
    if tblPr is None:
        tblPr = OxmlElement("w:tblPr"); tbl_el.insert(0, tblPr)

    # Set tblInd
    for old in tblPr.findall(qn("w:tblInd")):
        tblPr.remove(old)
    tblInd = OxmlElement("w:tblInd")
    tblInd.set(qn("w:w"), str(indent_dxa))
    tblInd.set(qn("w:type"), "dxa")
    tblPr.append(tblInd)

    # Update tblW
    tblW_el = tblPr.find(qn("w:tblW"))
    if tblW_el is None:
        tblW_el = OxmlElement("w:tblW"); tblPr.append(tblW_el)
    old_type = tblW_el.get(qn("w:type"), "dxa")
    old_w_str = tblW_el.get(qn("w:w"), "0")
    # Compute old width in dxa
    if old_type == "pct":
        old_w_dxa = int(old_w_str) * CONTENT_WIDTH // 5000
    else:
        old_w_dxa = int(old_w_str)
    tblW_el.set(qn("w:type"), "dxa")
    tblW_el.set(qn("w:w"), str(target_width))

    # Scale gridCol widths proportionally
    if old_w_dxa > 0:
        scale = target_width / old_w_dxa
        grid = tbl_el.find(qn("w:tblGrid"))
        if grid is not None:
            cols = grid.findall(qn("w:gridCol"))
            new_widths = [max(200, round(int(c.get(qn("w:w"), 0)) * scale)) for c in cols]
            # Adjust last column to fix rounding
            diff = target_width - sum(new_widths)
            new_widths[-1] += diff
            for c, w in zip(cols, new_widths):
                c.set(qn("w:w"), str(w))
            # Also update cell widths in all rows
            for tr in tbl_el.findall(qn("w:tr")):
                cells = tr.findall(qn("w:tc"))
                # Match cells to grid columns (handle colSpan)
                col_idx = 0
                for tc in cells:
                    tcPr = tc.find(qn("w:tcPr"))
                    if tcPr is not None:
                        tcW = tcPr.find(qn("w:tcW"))
                        gridSpan = tcPr.find(qn("w:gridSpan"))
                        span = int(gridSpan.get(qn("w:val"), 1)) if gridSpan is not None else 1
                        if tcW is not None and col_idx < len(new_widths):
                            cell_w = sum(new_widths[col_idx:col_idx+span])
                            tcW.set(qn("w:w"), str(cell_w))
                            tcW.set(qn("w:type"), "dxa")
                        col_idx += span


def make_section_heading_30(prefix_text, body_text):
    """Heading matching "| Перечень оборудования" style (styleId=30, Open Sans Condensed, red prefix)."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle"); pStyle.set(qn("w:val"), "30"); pPr.append(pStyle)
    sp = OxmlElement("w:spacing"); sp.set(qn("w:after"), "0"); pPr.append(sp)
    ind = OxmlElement("w:ind"); ind.set(qn("w:firstLine"), "284"); pPr.append(ind)
    rPr_p = OxmlElement("w:rPr")
    rf_p = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs"): rf_p.set(qn(a), "Open Sans Condensed")
    rPr_p.append(rf_p); pPr.append(rPr_p)
    p.append(pPr)
    r1 = OxmlElement("w:r")
    rPr1 = OxmlElement("w:rPr")
    rf1 = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs"): rf1.set(qn(a), "Open Sans Condensed")
    rPr1.append(rf1)
    clr1 = OxmlElement("w:color"); clr1.set(qn("w:val"), "E12222"); rPr1.append(clr1)
    r1.append(rPr1)
    t1 = OxmlElement("w:t"); t1.set(qn("xml:space"), "preserve"); t1.text = prefix_text; r1.append(t1)
    p.append(r1)
    r2 = OxmlElement("w:r")
    rPr2 = OxmlElement("w:rPr")
    rf2 = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs"): rf2.set(qn(a), "Open Sans Condensed")
    rPr2.append(rf2); r2.append(rPr2)
    t2 = OxmlElement("w:t"); t2.text = body_text; r2.append(t2)
    p.append(r2)
    return p


def get_text(el):
    return "".join(t.text or "" for t in el.iter(qn("w:t")))


def has_drawing(el):
    return len(list(el.iter(qn("w:drawing")))) > 0


def set_para_text(para_el, new_text):
    """Replace text runs keeping drawing runs and pPr intact."""
    runs = list(para_el.findall(qn("w:r")))
    drawing_runs = [r for r in runs if r.find(qn("w:drawing")) is not None]
    text_runs    = [r for r in runs if r.find(qn("w:drawing")) is None]
    first_rPr = None
    if text_runs:
        rp = text_runs[0].find(qn("w:rPr"))
        if rp is not None:
            first_rPr = copy.deepcopy(rp)
    for r in text_runs:
        para_el.remove(r)
    new_r = OxmlElement("w:r")
    if first_rPr is not None:
        new_r.append(first_rPr)
    new_t = OxmlElement("w:t")
    new_t.set(qn("xml:space"), "preserve")
    new_t.text = new_text
    new_r.append(new_t)
    para_el.append(new_r)


def set_cell_text(cell_el, new_text):
    paras = cell_el.findall(qn("w:p"))
    if paras:
        set_para_text(paras[0], new_text)
        for p in paras[1:]:
            cell_el.remove(p)


# ═══════════════════════════════════════════════════
# 1. REMOVE TECH PROCESS + EQUIPMENT DESCRIPTIONS
# ═══════════════════════════════════════════════════
children = list(body)
for i in sorted(list(range(21, 32)) + list(range(50, 165)), reverse=True):
    body.remove(children[i])
children = list(body)

# ═══════════════════════════════════════════════════
# 2. UPDATE TITLE TEXTS
# ═══════════════════════════════════════════════════
for child in children:
    text = get_text(child).strip()
    if text == "ЛИНИЯ":
        set_para_text(child, "ГРАНУЛЯТОР")
    elif text == "БРИКЕТИРОВАНИЯ":
        set_para_text(child, "ОГМ-1,5")          # keeps vertical bar drawing
    elif "Производительность" in text and "1 000" in text:
        set_para_text(child, "Производительность — 0,9–1,2 т/час")
    elif "| Перечень оборудования" in text:
        # Replace entire element with a properly styled heading
        new_heading = make_section_heading_30("| ", "Комплектация и стоимость")
        idx = list(body).index(child)
        body.remove(child)
        body.insert(idx, new_heading)
    elif "агрегат не является" in text:
        set_para_text(child, "")

# ═══════════════════════════════════════════════════
# 3. DETACH: params heading + params table + desc
#    (will be placed on page 2)
# ═══════════════════════════════════════════════════
children = list(body)

params_heading_idx = None
params_tbl_idx     = None
desc_idx           = None
desc_deco_idx      = None  # floating image paragraph just before description

for i, child in enumerate(children):
    tag  = child.tag.split("}")[-1]
    text = get_text(child)
    if "ПАРАМЕТРЫ ЛИНИИ" in text and params_heading_idx is None:
        params_heading_idx = i
    elif tag == "tbl" and params_heading_idx is not None and params_tbl_idx is None:
        params_tbl_idx = i
    elif "Разрабатываем и производим" in text and desc_idx is None:
        desc_idx = i
        # The decorative image paragraph is the one immediately before desc_idx
        # that has a drawing (search backwards)
        for j in range(i - 1, max(i - 5, 0), -1):
            if has_drawing(children[j]):
                desc_deco_idx = j
                break

print(f"params_heading={params_heading_idx}, params_tbl={params_tbl_idx}, desc_deco={desc_deco_idx}, desc={desc_idx}")

params_heading_el = children[params_heading_idx]
params_tbl_el     = children[params_tbl_idx]
desc_el           = children[desc_idx]
desc_deco_el      = children[desc_deco_idx] if desc_deco_idx is not None else None

body.remove(params_heading_el)
body.remove(params_tbl_el)
body.remove(desc_el)
if desc_deco_el is not None:
    body.remove(desc_deco_el)
children = list(body)

# ═══════════════════════════════════════════════════
# 4. REBUILD PARAMS TABLE FROM SCRATCH (clean style)
# ═══════════════════════════════════════════════════
NEW_PARAMS = [
    ("Производительность",        "0,9–1,2 т/ч"),
    ("Мощность двигателя",        "75–90 кВт"),
    ("Частота вращения матрицы",  "140 об/мин"),
    ("Диаметр отверстий матрицы", "6 мм / 8 мм"),
    ("Масса",                     "2,1 т"),
    ("Габариты (Д × Ш × В)",     "2 407 × 765 × 1 430 мм"),
    ("Питание",                   "380 В / 50 Гц"),
    ("Сырьё",                     "Опилки, стружка, солома, лузга"),
    ("Влажность сырья",           "10–15%"),
    ("Выходной продукт",          "Топливные пеллеты Ø 6/8 мм"),
]

# Build fresh table — same style as conditions table
p_col1, p_col2 = 3800, 6556  # total = 10356
params_tbl_el = OxmlElement("w:tbl")
ptblPr = OxmlElement("w:tblPr")
ptblW = OxmlElement("w:tblW")
ptblW.set(qn("w:w"), str(p_col1 + p_col2))
ptblW.set(qn("w:type"), "dxa")
ptblPr.append(ptblW)
params_tbl_el.append(ptblPr)
ptblGrid = OxmlElement("w:tblGrid")
for w in [p_col1, p_col2]:
    gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); ptblGrid.append(gc)
params_tbl_el.append(ptblGrid)

for label, value in NEW_PARAMS:
    tr = OxmlElement("w:tr")
    tr.append(make_cell(label, bold=True, width=p_col1, shading="F5F5F5"))
    tr.append(make_cell(value, bold=True, width=p_col2))
    params_tbl_el.append(tr)

apply_tbl_indent(params_tbl_el, 421)
print("Params table rebuilt (clean style)")

# ═══════════════════════════════════════════════════
# 5. UPDATE DESCRIPTION TEXT
# ═══════════════════════════════════════════════════
set_para_text(
    desc_el,
    "Промышленный пресс-гранулятор с кольцевой матрицей для производства "
    "топливных пеллет из опилок, стружки, соломы и лузги. "
    "Производительность до 1,2 т/ч. Матрица и ролики из легированной "
    "стали. Поставляется отдельно или в составе линии гранулирования под ключ.",
)

# ═══════════════════════════════════════════════════
# 6. INSERT IMAGE ON TITLE PAGE (where desc was)
# ═══════════════════════════════════════════════════
children = list(body)

# Find the empty area below divider line — between idx ~14 and equip heading
# Insert image after the red line area (around idx 15 — before empty paragraphs)
# Find the paragraph just before "| Комплектация"
equip_head_idx = None
for i, child in enumerate(children):
    if "| Комплектация" in get_text(child):
        equip_head_idx = i
        break

# Add image
temp_p = doc.add_paragraph()
temp_p.add_run().add_picture("ogm-image.png", width=Cm(13))
img_el = temp_p._element
pPr_img = OxmlElement("w:pPr")
jc = OxmlElement("w:jc"); jc.set(qn("w:val"), "center")
sp_img = OxmlElement("w:spacing"); sp_img.set(qn("w:before"), "300"); sp_img.set(qn("w:after"), "0")
pPr_img.append(jc); pPr_img.append(sp_img)
img_el.insert(0, pPr_img)
body.remove(img_el)

# Insert image just before equip heading (on title page, after divider line area)
body.insert(equip_head_idx, img_el)
print("Image inserted on title page")

# ═══════════════════════════════════════════════════
# 7. BUILD STYLED HEADING (style 30 — like original red headers)
# ═══════════════════════════════════════════════════
def make_section_heading_30(prefix_text, body_text):
    """Create heading matching "| Перечень оборудования" style (styleId=30)."""
    p = OxmlElement("w:p")
    pPr = OxmlElement("w:pPr")
    pStyle = OxmlElement("w:pStyle"); pStyle.set(qn("w:val"), "30"); pPr.append(pStyle)
    sp = OxmlElement("w:spacing"); sp.set(qn("w:after"), "0"); pPr.append(sp)
    ind = OxmlElement("w:ind"); ind.set(qn("w:firstLine"), "284"); pPr.append(ind)
    rPr_p = OxmlElement("w:rPr")
    rf_p = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs"): rf_p.set(qn(a), "Open Sans Condensed")
    rPr_p.append(rf_p); pPr.append(rPr_p)
    p.append(pPr)
    # "| " in red
    r1 = OxmlElement("w:r")
    rPr1 = OxmlElement("w:rPr")
    rf1 = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs"): rf1.set(qn(a), "Open Sans Condensed")
    rPr1.append(rf1)
    clr1 = OxmlElement("w:color"); clr1.set(qn("w:val"), "E12222"); rPr1.append(clr1)
    r1.append(rPr1)
    t1 = OxmlElement("w:t"); t1.set(qn("xml:space"), "preserve"); t1.text = prefix_text; r1.append(t1)
    p.append(r1)
    # body text
    r2 = OxmlElement("w:r")
    rPr2 = OxmlElement("w:rPr")
    rf2 = OxmlElement("w:rFonts")
    for a in ("w:ascii", "w:hAnsi", "w:cs"): rf2.set(qn(a), "Open Sans Condensed")
    rPr2.append(rf2); r2.append(rPr2)
    t2 = OxmlElement("w:t"); t2.text = body_text; r2.append(t2)
    p.append(r2)
    return p

tech_heading_el = make_section_heading_30("| ", "Технические характеристики")

# ═══════════════════════════════════════════════════
# 8. INSERT PAGE 2: desc → heading → params_tbl
#    before "| Комплектация"
# ═══════════════════════════════════════════════════
children = list(body)
equip_head_idx = None
for i, child in enumerate(children):
    if "| Комплектация" in get_text(child):
        equip_head_idx = i
        break

# Order: desc_deco_el, desc_el, spacer, tech_heading_el, params_tbl_el → then | Комплектация
insert_els = []
if desc_deco_el is not None:
    insert_els.append(desc_deco_el)
insert_els += [desc_el, make_para("", size=12), tech_heading_el, params_tbl_el]
for el in reversed(insert_els):
    body.insert(equip_head_idx, el)

print("Page 2: description + heading + params inserted")

# ═══════════════════════════════════════════════════
# 9. REBUILD EQUIPMENT TABLE
# ═══════════════════════════════════════════════════
children = list(body)
equip_tbl = None
equip_head_passed = False
for child in children:
    tag = child.tag.split("}")[-1]
    if "| Комплектація" in get_text(child) or "| Комплектация" in get_text(child):
        equip_head_passed = True
    if tag == "tbl" and equip_head_passed and equip_tbl is None:
        equip_tbl = child

if equip_tbl is not None:
    rows = equip_tbl.findall(qn("w:tr"))
    header_tmpl  = copy.deepcopy(rows[0])
    section_tmpl = copy.deepcopy(rows[1])
    data_tmpl    = copy.deepcopy(rows[2])
    for r in rows:
        equip_tbl.remove(r)

    def make_header_row(texts):
        row = copy.deepcopy(header_tmpl)
        cells = row.findall(qn("w:tc"))
        for i, cell in enumerate(cells):
            set_cell_text(cell, texts[i] if i < len(texts) else "")
        return row

    def make_section_row(text):
        row = copy.deepcopy(section_tmpl)
        cells = row.findall(qn("w:tc"))
        set_cell_text(cells[0], text)
        return row

    def make_data_row(num, name, price):
        row = copy.deepcopy(data_tmpl)
        cells = row.findall(qn("w:tc"))
        if len(cells) >= 4:
            set_cell_text(cells[0], num)
            set_cell_text(cells[1], name)
            set_cell_text(cells[2], "")
            set_cell_text(cells[3], price)
        return row

    equip_tbl.append(make_header_row(["№", "НАИМЕНОВАНИЕ", "", "СТОИМОСТЬ, РУБ."]))
    equip_tbl.append(make_section_row("СТАНДАРТНАЯ КОМПЛЕКТАЦИЯ"))
    equip_tbl.append(make_data_row("1", "Гранулятор ОГМ-1,5 с кольцевой матрицей d6 мм", "3 014 000"))
    equip_tbl.append(make_data_row("2", "Гранулятор ОГМ-1,5 с кольцевой матрицей d8 мм", "3 014 000"))
    equip_tbl.append(make_section_row("ДОПОЛНИТЕЛЬНЫЕ ОПЦИИ"))
    equip_tbl.append(make_data_row("3", "Частотный преобразователь (мягкий пуск, защита двигателя)", "по запросу"))
    equip_tbl.append(make_data_row("4", "Система автоматической смазки прессующего узла", "по запросу"))
    equip_tbl.append(make_section_row("РАСХОДНЫЕ МАТЕРИАЛЫ"))
    equip_tbl.append(make_data_row("5", "Матрица кольцевая d6 мм (легированная сталь)", "по запросу"))
    equip_tbl.append(make_data_row("6", "Матрица кольцевая d8 мм (легированная сталь)", "по запросу"))
    equip_tbl.append(make_data_row("7", "Ролики прессующие (комплект)", "по запросу"))
    apply_tbl_indent(equip_tbl, 421)
    print("Equipment table rebuilt")

# ═══════════════════════════════════════════════════
# 10. REMOVE FANCY "НАШИ УСЛУГИ" + INSERT CONDITIONS/CONTACTS
# ═══════════════════════════════════════════════════
children = list(body)
remove_start = None
sectPr_idx   = None
for i, child in enumerate(children):
    if "Наши услуги" in get_text(child) and remove_start is None:
        remove_start = i
    if child.tag.split("}")[-1] == "sectPr":
        sectPr_idx = i
        break

for i in range(remove_start, sectPr_idx)[::-1]:
    body.remove(children[i])
print(f"Removed services section ({remove_start}..{sectPr_idx-1})")


# Page break
pb = OxmlElement("w:p")
pb_r = OxmlElement("w:r"); br = OxmlElement("w:br"); br.set(qn("w:type"), "page"); pb_r.append(br); pb.append(pb_r)

h_cond = make_section_heading_30("| ", "Условия поставки")

# content=11198, tblInd=421 left + 421 right → table width = 11198 - 421 - 421 = 10356
col1, col2 = 3200, 7156  # 3200 + 7156 = 10356
conditions = [
    ("Шеф-монтаж",        "20 000 руб/день + проживание сотрудников на территории Заказчика"),
    ("Полный монтаж",     "Установка под ключ. Стоимость — 15% от стоимости оборудования"),
    ("Место передачи",    "Склад Поставщика, г. Белая Холуница, Кировская область"),
    ("Доставка",          "Поставщик организует доставку за счёт Покупателя"),
    ("Срок изготовления", "60 рабочих дней с момента внесения аванса"),
    ("Условия оплаты",    "30% аванс, 70% по уведомлению о готовности"),
    ("Гарантия",          "36 месяцев с момента получения оборудования"),
]
cond_tbl = OxmlElement("w:tbl")
tP = OxmlElement("w:tblPr"); tW = OxmlElement("w:tblW"); tW.set(qn("w:w"), str(col1 + col2)); tW.set(qn("w:type"), "dxa"); tP.append(tW); cond_tbl.append(tP)
tG = OxmlElement("w:tblGrid")
for w in [col1, col2]:
    gc = OxmlElement("w:gridCol"); gc.set(qn("w:w"), str(w)); tG.append(gc)
cond_tbl.append(tG)
for label, value in conditions:
    tr = OxmlElement("w:tr")
    tr.append(make_cell(label, bold=True, width=col1, shading="F5F5F5"))
    tr.append(make_cell(value, width=col2))
    cond_tbl.append(tr)
apply_tbl_indent(cond_tbl, 421)

h_cont = make_section_heading_30("| ", "Контакты")

CONT_INDENT = 421  # match table left indent

p_phone = OxmlElement("w:p")
ph_pPr = OxmlElement("w:pPr")
ph_sp = OxmlElement("w:spacing"); ph_sp.set(qn("w:after"), "80"); ph_pPr.append(ph_sp)
ph_ind = OxmlElement("w:ind"); ph_ind.set(qn("w:left"), str(CONT_INDENT)); ph_pPr.append(ph_ind)
p_phone.append(ph_pPr)
p_phone.append(make_run("Тел.: ", bold=False, size=22, color="4B5464"))
p_phone.append(make_run("+7 900 521-84-77", bold=True, size=22))

sig2 = OxmlElement("w:p")
sig2_pPr = OxmlElement("w:pPr")
sig2_sp = OxmlElement("w:spacing"); sig2_sp.set(qn("w:after"), "100"); sig2_pPr.append(sig2_sp)
sig2_ind = OxmlElement("w:ind"); sig2_ind.set(qn("w:left"), str(CONT_INDENT)); sig2_pPr.append(sig2_ind)
sig2.append(sig2_pPr)
sig2.append(make_run("Отдел продаж ", size=22))
sig2.append(make_run('ООО "КировБелМаш"', bold=True, size=22))

contacts = [
    make_para('ООО "КировБелМаш"', bold=True, size=24, space_after=80, indent_left=CONT_INDENT),
    make_para("Кировская область, г. Белая Холуница, ул. Глазырина, 112", size=22, space_after=80, indent_left=CONT_INDENT),
    p_phone,
    make_para("Email: sale@kirovbelmash.ru", size=22, space_after=80, indent_left=CONT_INDENT),
    make_para("Сайт: kirovbelmash.ru", size=22, space_after=80, indent_left=CONT_INDENT),
]

children = list(body)
sectPr = next(c for c in children if c.tag.split("}")[-1] == "sectPr")
for el in [pb, h_cond, cond_tbl, make_para("", size=12), h_cont, *contacts,
           make_para("", size=12),
           make_para("С уважением,", size=22, color="4B5464", space_before=400, indent_left=CONT_INDENT),
           sig2]:
    body.insert(list(body).index(sectPr), el)


# ═══════════════════════════════════════════════════
# FINAL CHECK
# ═══════════════════════════════════════════════════
children = list(body)
print(f"\nFinal: {len(children)} elements")
for i, child in enumerate(children):
    tag  = child.tag.split("}")[-1]
    text = get_text(child)[:65]
    img  = " [IMG]" if has_drawing(child) else ""
    if text or tag in ("tbl", "sectPr"):
        print(f"  {i}: <{tag}>{img} '{text}'")

doc.save("КП_Гранулятор_ОГМ-1,5_v3.docx")
print("\nSaved: КП_Гранулятор_ОГМ-1,5_v3.docx")
